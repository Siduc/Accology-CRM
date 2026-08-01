"""
Build Accology Limited letterhead as Word (.docx).

Layout matches Grok Imagine post (modern left-aligned lockup, cyan accent bar,
thin rule, contact strip). Legal/contact lines use real Accology Limited details
(Companies House 07210650).
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor, Twips
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(__file__).resolve().parents[1]
BRAND = ROOT / "app" / "static" / "branding"
OUT = BRAND / "letterhead.docx"
HEADER_PNG = BRAND / "letterhead_header.png"
REF_IMG = BRAND / "landing_imagine_letterhead.jpg"

# Brand (landing + Imagine)
CYAN = (0, 229, 255)  # #00E5FF
PURPLE = (124, 58, 237)  # gradient end
INK = (15, 23, 42)  # near-black
AS_TEAL = (30, 64, 100)
MUTED = (100, 110, 120)
LINE = (0, 200, 220)

# Accology Limited (CH 07210650)
COMPANY_NO = "07210650"
REG_OFFICE = "USN Bolton Arena, Arena Approach, Horwich, Bolton, BL6 6LB"
EMAIL = "simon@accology.co"
MOBILE = "07857 224801"
TEL = "01204 238938"
FAX = "01204 238939"
VAT = "203313763"
AUDITOR = "2143851"
DIRECTOR = "Simon Duckworth – Director"


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = []
    if bold:
        candidates += [
            r"C:\Windows\Fonts\arialbd.ttf",
            r"C:\Windows\Fonts\segoeuib.ttf",
            r"C:\Windows\Fonts\calibrib.ttf",
        ]
    candidates += [
        r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\segoeui.ttf",
        r"C:\Windows\Fonts\calibri.ttf",
    ]
    for p in candidates:
        if Path(p).is_file():
            try:
                return ImageFont.truetype(p, size)
            except OSError:
                continue
    return ImageFont.load_default()


def build_header_png(path: Path, width_px: int = 1650, height_px: int = 320) -> Path:
    """Header band matching Imagine: ACCOLOGY + gradient bar + Accounting Science + rule + contact."""
    img = Image.new("RGB", (width_px, height_px), (255, 255, 255))
    draw = ImageDraw.Draw(img)

    font_word = _font(72, bold=True)
    font_tag = _font(36, bold=False)
    font_contact = _font(22, bold=False)

    x0, y0 = 8, 18
    word = "ACCOLOGY"
    draw.text((x0, y0), word, fill=INK, font=font_word)
    bbox = draw.textbbox((x0, y0), word, font=font_word)
    word_w = bbox[2] - bbox[0]
    bar_y = bbox[3] + 6
    bar_h = 8
    # Cyan → purple gradient bar under ACCOLOGY
    for i in range(word_w):
        t = i / max(word_w - 1, 1)
        r = int(CYAN[0] + (PURPLE[0] - CYAN[0]) * t)
        g = int(CYAN[1] + (PURPLE[1] - CYAN[1]) * t)
        b = int(CYAN[2] + (PURPLE[2] - CYAN[2]) * t)
        draw.line([(x0 + i, bar_y), (x0 + i, bar_y + bar_h)], fill=(r, g, b))

    tag = "Accounting Science"
    tag_y = bar_y + bar_h + 14
    draw.text((x0, tag_y), tag, fill=AS_TEAL, font=font_tag)
    tag_bbox = draw.textbbox((x0, tag_y), tag, font=font_tag)

    # Thin cyan full-width rule
    rule_y = tag_bbox[3] + 18
    draw.line([(x0, rule_y), (width_px - 8, rule_y)], fill=LINE, width=2)

    contact = (
        f"{REG_OFFICE}  |  Tel: {TEL}  |  Mobile: {MOBILE}  |  {EMAIL}"
    )
    draw.text((x0, rule_y + 14), contact, fill=MUTED, font=font_contact)

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, "PNG", optimize=True)
    return path


def _set_run_font(run, size=8, color=(0x64, 0x6E, 0x78), bold=False, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)


def _set_paragraph_border_top(paragraph, color="00C8DC", sz="12"):
    """Thin top border on a footer paragraph (Imagine-style rule)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), sz)
    top.set(qn("w:space"), "8")
    top.set(qn("w:color"), color)
    pBdr.append(top)
    pPr.append(pBdr)


def build() -> Path:
    BRAND.mkdir(parents=True, exist_ok=True)
    build_header_png(HEADER_PNG)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(28)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(10)

    # —— Header image (Imagine-style lockup + real contact strip) ——
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_before = Pt(0)
    hp.paragraph_format.space_after = Pt(0)
    run = hp.add_run()
    # Content width ≈ 210 − 22 − 22 = 166 mm
    run.add_picture(str(HEADER_PNG), width=Mm(166))

    # —— Footer (real Accology Limited legal lines) ——
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(6)
    fp.paragraph_format.space_after = Pt(2)
    _set_paragraph_border_top(fp, color="CBD5E1", sz="6")

    lines = [
        f"{DIRECTOR}  ·  Email: {EMAIL}  ·  Mobile: {MOBILE}  ·  Tel: {TEL}",
        f"Accology is a trading name of Accology Limited – Company number {COMPANY_NO}",
        f"Registered office – {REG_OFFICE}",
        f"Accology Limited is a registered auditor authorised by The Association of Certified Accountants – Registered number {AUDITOR}",
        f"VAT No. {VAT}",
    ]
    first = True
    for line in lines:
        if first:
            p = fp
            first = False
        else:
            p = footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
        r = p.add_run(line)
        _set_run_font(r, size=7.5, color=(0x64, 0x6E, 0x78))

    # —— Body canvas ——
    body = doc.add_paragraph("")
    body.paragraph_format.space_before = Pt(12)
    for _ in range(5):
        doc.add_paragraph("")

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path} ({path.stat().st_size} bytes)")
    print(f"Header PNG: {HEADER_PNG} ({HEADER_PNG.stat().st_size} bytes)")
    if REF_IMG.is_file():
        print(f"Imagine reference: {REF_IMG.name}")
